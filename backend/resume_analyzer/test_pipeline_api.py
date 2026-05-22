import requests
import sys
import os
from docx import Document

def create_valid_docx(filename):
    doc = Document()
    doc.add_heading("RAHUL SHARMA", 0)
    doc.add_paragraph("rahul.sharma@email.com | +91-9876543210 | Bangalore, India")
    doc.add_heading("EDUCATION", level=1)
    doc.add_paragraph("B.Tech in Computer Science and Engineering | CGPA: 8.7/10")
    doc.add_heading("TECHNICAL SKILLS", level=1)
    doc.add_paragraph("Python, Java, JavaScript, SQL, React.js, Node.js, AWS, Docker, Machine Learning")
    doc.add_heading("PROJECTS", level=1)
    doc.add_paragraph("AI-Powered Resume Analyzer - Developed an intelligent resume parsing system using NLP.")
    doc.save(filename)

def run_test():
    base_url = "http://localhost:8000"
    docx_filename = "test_resume.docx"
    create_valid_docx(docx_filename)

    print("Sending ping/health check...")
    try:
        r = requests.get(f"{base_url}/health")
        print(f"Health response: {r.status_code} - {r.json()}")
    except Exception as e:
        print(f"Could not connect to backend: {e}")
        if os.path.exists(docx_filename):
            os.remove(docx_filename)
        sys.exit(1)

    print("\n--- STAGE 1: Uploading Resume ---")
    try:
        with open(docx_filename, "rb") as f:
            files = {"file": (docx_filename, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            r = requests.post(f"{base_url}/api/analyze/file", files=files)
    finally:
        if os.path.exists(docx_filename):
            os.remove(docx_filename)

    print(f"Upload Status Code: {r.status_code}")
    try:
        resume_data = r.json()
        print(f"Upload Response JSON: {resume_data}")
    except Exception as e:
        print(f"Failed to parse upload JSON: {e}")
        return

    if r.status_code != 200:
        print("Upload failed. Stopping test.")
        return

    # Extract skills
    skills = resume_data.get("analysis", {}).get("skills", {}).get("found_skills", [])
    print(f"Skills Extracted: {skills}")

    profile = {
        "cgpa": None,
        "has_backlogs": False,
        "branch": "CSE",
        "top_n": 10,
        "mock_test_score": None
    }
    print(f"Profile: {profile}")

    print("\n--- STAGE 2: Get Job Recommendations ---")
    payload_jobs = {
        "skills": skills,
        "cgpa": 7.0,
        "branch": "CSE",
        "has_backlogs": False,
        "top_n": 10
    }
    print(f"Jobs Payload: {payload_jobs}")
    r = requests.post(f"{base_url}/api/recommend-jobs", json=payload_jobs)
    print(f"Jobs Response Status: {r.status_code}")
    try:
        job_data = r.json()
        print(f"Jobs Response JSON (truncated): {str(job_data)[:200]}...")
    except Exception as e:
        print(f"Failed to parse jobs JSON: {e}")
        return

    if r.status_code != 200:
        print("Jobs step failed.")
        return

    print("\n--- STAGE 3: Predict Placement ---")
    summary = resume_data.get("summary", {})
    category_scores = resume_data.get("analysis", {}).get("scoring", {}).get("category_scores", {})
    top_job = job_data.get("top_recommendations", [{}])[0] if job_data.get("top_recommendations") else {}
    
    payload_predict = {
        "cgpa": 7.0,
        "resume_score": summary.get("overall_score", 50),
        "skills_count": summary.get("total_skills", 0),
        "technical_skills": summary.get("technical_skills", 0),
        "soft_skills": summary.get("soft_skills", 0),
        "high_demand_skills": summary.get("high_demand_skills", 0),
        "projects_count": summary.get("projects_count", 0),
        "exp_months": summary.get("experience_count", 0) * 3,
        "certifications": summary.get("certifications_count", 0),
        "job_match_score": top_job.get("match_score", 0),
        "ats_score": category_scores.get("ats_compatibility", 50),
        "skill_diversity_score": resume_data.get("analysis", {}).get("skills", {}).get("skill_diversity_score", 50),
        "has_backlogs": False,
        "branch": "CSE",
        "model_choice": "ensemble"
    }
    print(f"Predict Payload: {payload_predict}")
    r = requests.post(f"{base_url}/api/predict", json=payload_predict)
    print(f"Predict Response Status: {r.status_code}")
    try:
        predict_data = r.json()
        print(f"Predict Response JSON: {predict_data}")
    except Exception as e:
        print(f"Failed to parse predict JSON: {e}")
        return

    if r.status_code != 200:
        print("Predict step failed.")
        return

    print("\n--- STAGE 4: Placement Readiness ---")
    payload_readiness = {
        "cgpa": 7.0,
        "branch": "CSE",
        "has_backlogs": False,
        "resume_score": summary.get("overall_score", 50),
        "total_skills": summary.get("total_skills", 0),
        "technical_skills": summary.get("technical_skills", 0),
        "high_demand_skills": summary.get("high_demand_skills", 0),
        "skill_diversity": resume_data.get("analysis", {}).get("skills", {}).get("skill_diversity_score", 50),
        "projects_count": summary.get("projects_count", 0),
        "certifications": summary.get("certifications_count", 0),
        "ats_score": category_scores.get("ats_compatibility", 50),
        "job_match_score": top_job.get("match_score", 0),
        "total_jobs_matched": job_data.get("total_jobs_matched", 0),
        "selection_probability": predict_data.get("selection_probability", 50),
        "mock_test_score": None
    }
    print(f"Readiness Payload: {payload_readiness}")
    r = requests.post(f"{base_url}/api/readiness/direct", json=payload_readiness)
    print(f"Readiness Response Status: {r.status_code}")
    try:
        readiness_data = r.json()
        print(f"Readiness Response JSON: {readiness_data}")
    except Exception as e:
        print(f"Failed to parse readiness JSON: {e}")
        return

    print("\nAll stages simulated.")

if __name__ == "__main__":
    run_test()
