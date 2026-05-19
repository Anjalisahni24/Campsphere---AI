"""
CAMSPHER-AI Resume Analyzer
Text Extraction Module - Extracts text from PDF and DOCX resumes
<<<<<<< HEAD
=======
<<<<<<< HEAD
=======

v1.1.0 — Added pdfminer.six as third fallback for scanned/encoded PDFs
>>>>>>> 66e74765270250cf239d6dba7e73fe97b971a73a
>>>>>>> fe183888c2042d6c21e43802e39f44db90f765b1
"""

import io
import re
from typing import Optional

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
<<<<<<< HEAD
=======
<<<<<<< HEAD
=======
    from pdfminer.high_level import extract_text as pdfminer_extract
    from pdfminer.layout import LAParams
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False

try:
>>>>>>> 66e74765270250cf239d6dba7e73fe97b971a73a
>>>>>>> fe183888c2042d6c21e43802e39f44db90f765b1
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


class TextExtractor:
    """Extracts and cleans text from resume files (PDF, DOCX)."""

    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc']

    def extract(self, file_content: bytes, filename: str) -> str:
        """
        Extract text from file content based on file extension.

        Args:
            file_content: Raw bytes of the file
            filename: Name of the file with extension

        Returns:
            Cleaned extracted text
        """
        filename_lower = filename.lower()

        if filename_lower.endswith('.pdf'):
            return self._extract_pdf(file_content)
        elif filename_lower.endswith(('.docx', '.doc')):
            return self._extract_docx(file_content)
        else:
            raise ValueError(f"Unsupported file format. Supported: {self.supported_formats}")

    def _extract_pdf(self, file_content: bytes) -> str:
<<<<<<< HEAD
=======
<<<<<<< HEAD
>>>>>>> fe183888c2042d6c21e43802e39f44db90f765b1
        """Extract text from PDF using pdfplumber (preferred) or PyPDF2."""
        text = ""

        # Try pdfplumber first (better extraction quality)
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"
            except Exception:
                text = ""  # Reset and try fallback

        # Fallback to PyPDF2
        if not text.strip() and PYPDF2_AVAILABLE:
            try:
                reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            except Exception:
                pass

        if not text.strip():
            raise RuntimeError("Failed to extract text from PDF. Please ensure the PDF is not scanned/image-based.")
<<<<<<< HEAD
=======
=======
        """
        Extract text from PDF using 3 methods in order of quality:
          1. pdfplumber  — best for structured/table PDFs
          2. PyPDF2      — fallback for simple PDFs
          3. pdfminer.six — handles encoded/complex fonts (most compatible)
        """
        text = ""
        errors = []

        # ── Method 1: pdfplumber ─────────────────────────────────────────────
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    pages = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            pages.append(page_text)
                    if pages:
                        text = "\n\n".join(pages)
            except Exception as e:
                errors.append(f"pdfplumber: {e}")
                text = ""

        # ── Method 2: PyPDF2 ─────────────────────────────────────────────────
        if not text.strip() and PYPDF2_AVAILABLE:
            try:
                reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                pages = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        pages.append(page_text)
                if pages:
                    text = "\n\n".join(pages)
            except Exception as e:
                errors.append(f"PyPDF2: {e}")
                text = ""

        # ── Method 3: pdfminer.six ───────────────────────────────────────────
        # Handles PDFs with embedded fonts, unusual encoding, mixed layouts
        if not text.strip() and PDFMINER_AVAILABLE:
            try:
                laparams = LAParams(
                    line_margin=0.5,
                    word_margin=0.1,
                    char_margin=2.0,
                    boxes_flow=0.5,
                    detect_vertical=False,
                )
                extracted = pdfminer_extract(
                    io.BytesIO(file_content),
                    laparams=laparams,
                )
                if extracted and extracted.strip():
                    text = extracted
            except Exception as e:
                errors.append(f"pdfminer: {e}")
                text = ""

        # ── All methods failed ───────────────────────────────────────────────
        if not text.strip():
            err_detail = " | ".join(errors) if errors else "No text found"
            raise RuntimeError(
                f"Could not extract text from this PDF. "
                f"This usually means the PDF is scanned (image-only). "
                f"Please use the text input option instead — copy-paste your resume text directly. "
                f"[Technical: {err_detail}]"
            )
>>>>>>> 66e74765270250cf239d6dba7e73fe97b971a73a
>>>>>>> fe183888c2042d6c21e43802e39f44db90f765b1

        return self._clean_text(text)

    def _extract_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        if not PYTHON_DOCX_AVAILABLE:
<<<<<<< HEAD
=======
<<<<<<< HEAD
>>>>>>> fe183888c2042d6c21e43802e39f44db90f765b1
            raise RuntimeError("python-docx not installed. Install with: pip install python-docx")

        try:
            doc = Document(io.BytesIO(file_content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
<<<<<<< HEAD
=======
=======
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")

        try:
            doc = Document(io.BytesIO(file_content))

            # Extract from paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # Also extract from tables (many resumes use tables for layout)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text and cell_text not in paragraphs:
                            paragraphs.append(cell_text)

>>>>>>> 66e74765270250cf239d6dba7e73fe97b971a73a
>>>>>>> fe183888c2042d6c21e43802e39f44db90f765b1
            text = "\n\n".join(paragraphs)
            return self._clean_text(text)
        except Exception as e:
            raise RuntimeError(f"Failed to extract text from DOCX: {str(e)}")

    def _clean_text(self, text: str) -> str:
        """Clean extracted text by removing extra whitespace, special chars, etc."""
        # Remove null bytes
        text = text.replace('\x00', '')
        # Replace multiple newlines with double newline
        text = re.sub(r'\n{3,}', '\n\n', text)
        # Replace multiple spaces with single space
        text = re.sub(r' {2,}', ' ', text)
<<<<<<< HEAD
        # Remove non-printable characters except newlines
=======
<<<<<<< HEAD
        # Remove non-printable characters except newlines
=======
        # Remove non-printable characters except newlines and tabs
>>>>>>> 66e74765270250cf239d6dba7e73fe97b971a73a
>>>>>>> fe183888c2042d6c21e43802e39f44db90f765b1
        text = ''.join(char for char in text if char.isprintable() or char in '\n\t')
        # Strip whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
<<<<<<< HEAD
=======
<<<<<<< HEAD
=======
        # Remove lines that are just whitespace or single chars (noise)
        lines = [line for line in text.split('\n') if len(line.strip()) > 1 or line == '']
        text = '\n'.join(lines)
>>>>>>> 66e74765270250cf239d6dba7e73fe97b971a73a
>>>>>>> fe183888c2042d6c21e43802e39f44db90f765b1
        # Final strip
        return text.strip()

    def extract_from_text(self, text: str) -> str:
        """Clean raw pasted text (for direct text input)."""
        return self._clean_text(text)


def extract_text(file_content: bytes, filename: str) -> str:
    """Convenience function for text extraction."""
    extractor = TextExtractor()
<<<<<<< HEAD
    return extractor.extract(file_content, filename)
=======
<<<<<<< HEAD
    return extractor.extract(file_content, filename)
=======
    return extractor.extract(file_content, filename)
>>>>>>> 66e74765270250cf239d6dba7e73fe97b971a73a
>>>>>>> fe183888c2042d6c21e43802e39f44db90f765b1
